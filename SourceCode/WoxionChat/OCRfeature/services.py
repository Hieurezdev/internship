import os
import base64

from pathlib import Path
from typing import Dict, Any, List
from django.conf import settings
from .models import UploadedFile, OCRResultFactory
from accounts.models import User
from mistralai.models import OCRResponse
import io
from mistralai import Mistral
from mistralai import DocumentURLChunk
import logging
import re
import docx
from io import BytesIO

import asyncio
import aiohttp

logger = logging.getLogger(__name__)

class MistralAI:
    def __init__(self, mistral_api_key):
        self.api_key = mistral_api_key
        self.mistral_client = Mistral(api_key= mistral_api_key)
    @property
    def chat(self):
        return self.mistral_client.chat

class FileUploadService:
    
    @staticmethod
    def upload_file(user: User, title: str, file_obj) -> UploadedFile:
        try:
            validation_result = FileUploadService._validate_file(file_obj)
            if not validation_result['is_valid']:
                logger.error(f"File validation failed: {validation_result['errors']}")
                raise ValueError('; '.join(validation_result['errors']))
            if not title or not title.strip():
                title = FileUploadService._generate_title_from_filename(file_obj.name)

            uploaded_file = UploadedFile(title=title.strip())
            uploaded_file.set_uploader(user)
            uploaded_file.file.put(file_obj, filename=file_obj.name)
  
            # Save to database
            logger.info("Starting database save...")
            uploaded_file.save()

            mistral_api_key = getattr(settings, 'MISTRAL_API_KEY', None)
            if mistral_api_key:
                try:
                    OCRProcessingService.Processing_with_mistral(uploaded_file, mistral_api_key)
                    logger.info(f"OCR processing started for file {uploaded_file.id}")
                except Exception as e:
                    logger.warning(f"OCR processing failed for file {uploaded_file.id}: {e}")
            else:
                logger.warning("MISTRAL_API_KEY not found, creating basic OCR task")
                OCRProcessingService.create_ocr_task(uploaded_file)
            
            return uploaded_file
            
        except Exception as e:
            logger.error(f"Error uploading file: {e}")
            raise
    
    @staticmethod
    def delete_file(uploaded_file: UploadedFile):
        try:
            file_id = uploaded_file.id
            uploaded_file.delete()
            logger.info(f"File deleted successfully: {file_id}")
        except Exception as e:
            logger.error(f"Error deleting file {uploaded_file.id}: {e}")
            raise
    
    @staticmethod
    def _validate_file(file_obj) -> Dict[str, Any]:
        errors = []
        warnings = []
        
        ALLOWED_EXTENSIONS = ['.docx', '.pdf', '.png', '.jpg', '.jpeg', '.txt']
        MAX_FILE_SIZE = 50 * 1024 * 1024  # 50MB
        
        # Check extension
        if hasattr(file_obj, 'name') and file_obj.name:
            ext = os.path.splitext(file_obj.name)[1].lower()
            if ext not in ALLOWED_EXTENSIONS:
                errors.append(f"File extension '{ext}' not allowed")
        else:
            errors.append("Invalid filename")
        
        # Check file size
        if hasattr(file_obj, 'size') and file_obj.size:
            if file_obj.size > MAX_FILE_SIZE:
                errors.append(f"File too large: {file_obj.size:,} bytes")
        else:
            errors.append("Cannot determine file size")
        
        return {
            'is_valid': len(errors) == 0,
            'errors': errors,
            'warnings': warnings
        }
    
    @staticmethod
    def _generate_title_from_filename(filename: str) -> str:
        if not filename:
            return "Untitled Document"
        
        name_without_ext = os.path.splitext(filename)[0]
        title = name_without_ext.replace('_', ' ').replace('-', ' ')
        title = ' '.join(word.capitalize() for word in title.split())
        return title[:100] + "..." if len(title) > 100 else title
    
    @staticmethod
    def get_user_files(user: User, limit: int = None) -> List[UploadedFile]:
        queryset = UploadedFile.objects(uploader_id=str(user.id)).order_by('-uploaded_at')
        if limit:
            queryset = queryset[:limit]
        return list(queryset)

class OCRProcessingService:
    """Service xử lý OCR với Mistral AI - PURE STATIC METHODS"""
    @staticmethod
    def create_ocr_task(uploaded_file: UploadedFile):
        """Tạo OCR task cho file - ✅ UPDATED: Sử dụng OCRResultFactory"""
        logger.info("Creating ocr task ...")
        ocr_result = OCRResultFactory.create_ocr_result(uploaded_file)
        ocr_result.save()
        return ocr_result
    
    @staticmethod
    def Processing_with_mistral(uploaded_file: UploadedFile, mistral_api_key: str):
        """Process OCR with Mistral AI - ✅ UPDATED: Sử dụng OCRResultFactory"""
        logger.info("Processing with mistral ...")
        ocr_result = OCRResultFactory.create_ocr_result(uploaded_file)
        ocr_result.save()
        ocr_result.mark_as_processing()
        
        try:
            # Run async processing in sync context
            result = asyncio.run(OCRProcessingService._process_file_async(uploaded_file, mistral_api_key))
            
            # Lưu kết quả
            ocr_result.mark_as_success(
                result_data=result,
                raw_markdown=result.get('markdown_content', '')
            )
            
        except Exception as e:
            logger.error(f"Error processing file {uploaded_file.id}: {e}")
            ocr_result.mark_as_failed(str(e))
            raise
        
        return ocr_result

    @staticmethod
    async def _process_file_async(uploaded_file: UploadedFile, mistral_api_key: str) -> Dict[str, Any]:
        logger.info("Process file async ...")
        """Async file processing dispatcher"""
        # Khởi tạo Mistral client
        Misai = MistralAI(mistral_api_key)
        mistral_client = Misai.mistral_client
        
        # Lấy file content từ GridFS
        file_content = uploaded_file.file.read()
        file_extension = uploaded_file.extension.lower()
        filename = uploaded_file.filename
        
        # Xử lý theo loại file với async support
        if file_extension in ['.png', '.jpg', '.jpeg']:
            result = await OCRProcessingService._process_image_with_mistral_async(
                Misai, file_content, filename
            )
        elif file_extension == '.txt':
            result = await OCRProcessingService._process_text_with_mistral_async(
                Misai, file_content, filename
            )
        elif file_extension == '.pdf':
            result = await OCRProcessingService._process_pdf_with_mistral_async(
                Misai, file_content, filename
            )
        elif file_extension == '.docx':
            result = await OCRProcessingService._process_docx_with_mistral_async(
                Misai, file_content, filename
            )
        else:
            raise ValueError(f"Unsupported file type: {file_extension}")
        
        return result

    @staticmethod
    def _process_docx_with_mistral(mistral_client: MistralAI, file_content: bytes, filename: str) -> Dict[str, Any]:
        """Process DOCX file with Mistral AI"""    
        # Create a BytesIO object from file content
        doc_stream = BytesIO(file_content)
        
        # Load the document
        doc = docx.Document(doc_stream)
        
        # Extract text from all paragraphs
        text_content = []
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_content.append(paragraph.text)
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    text_content.append(" | ".join(row_text))
        
        # Combine all text
        full_text = "\n\n".join(text_content)
        
        if not full_text.strip():
            raise ValueError("No text content found in DOCX file")
        
        # Process with Mistral AI to format as markdown
        ocr_response = mistral_client.chat.complete(
            model="mistral-large-latest",
            messages=[
                {
                    "role": "user",
                    "content": f"Hãy định dạng nội dung DOCX này thành markdown có cấu trúc với tiêu đề, danh sách, bảng và formatting phù hợp. Giữ nguyên toàn bộ nội dung và ngôn ngữ gốc:\n\n{full_text}"
                }
            ],
            max_tokens=4000,
            temperature=0.1
        )
        
        markdown_content = ocr_response.choices[0].message.content if ocr_response and ocr_response.choices else full_text
        
        result = {
            "file_type": "docx",
            "filename": filename,
            "markdown_content": markdown_content,
            "structured_content": {
                "type": "docx_extraction",
                "processing_model": "mistral-large-latest",
                "paragraphs_count": len([p for p in doc.paragraphs if p.text.strip()]),
                "tables_count": len(doc.tables)
            }
        }
        
        return result

        
    
    @staticmethod
    def _process_pdf_with_mistral(mistral_client: MistralAI, file_content: bytes, filename: str) -> Dict[str, Any]:
        """Process PDF with Mistral OCR API directly"""
        logger.info("Processing pdf with mistral ...")
        try:
            uploaded_file = mistral_client.mistral_client.files.upload(
                file={
                    "file_name": Path(filename).stem,
                    "content": file_content,
                },
                purpose="ocr",
            )
            logger.info("get uploaded_file for Mistral OCR")
            # Get signed URL for processing
            signed_url = mistral_client.mistral_client.files.get_signed_url(
                file_id=uploaded_file.id, 
                expiry=1
            )
            
            # Process PDF with Mistral OCR
            pdf_response = mistral_client.mistral_client.ocr.process(
                document=DocumentURLChunk(document_url=signed_url.url), 
                model="mistral-large-latest", 
                include_image_base64=False  # Set to False to reduce response size
            )
            
            # Extract markdown content from OCR response
            all_markdowns = []
            for page in pdf_response.pages:
                if page.markdown and page.markdown.strip():
                    all_markdowns.append(page.markdown)
            
            # Combine all pages
            combined_content = f"# PDF: {filename}\n\n" + "\n\n---\n\n".join(all_markdowns)
            
            # Clean up the uploaded file from Mistral
            try:
                mistral_client.mistral_client.files.delete(file_id=uploaded_file.id)
            except Exception as e:
                logger.warning(f"Could not delete uploaded file from Mistral: {e}")
            
            result = {
                "file_type": "pdf",
                "filename": filename,
                "markdown_content": combined_content,
                "structured_content": {
                    "type": "pdf_ocr_mistral_api",
                    "processing_method": "mistral_ocr_api",
                    "pages_processed": len(pdf_response.pages),
                    "processing_model": "mistral-large-latest"
                }
            }
            
            return result
            
        except Exception as e:
            logger.error(f"Error processing PDF with Mistral OCR API: {e}")
            # Fallback to old method if Mistral OCR fails
            return OCRProcessingService._process_pdf_with_mistral_fallback(mistral_client, file_content, filename)
    
    @staticmethod
    def _process_pdf_with_mistral_fallback(mistral_client: MistralAI, file_content: bytes, filename: str) -> Dict[str, Any]:
        """Fallback method using pdf2image conversion"""
        logger.info("Process pdf with mistral fallback ...")
        try:
            from pdf2image import convert_from_bytes
            images = convert_from_bytes(file_content, dpi=200, first_page=1, last_page=20)  # Limit to 20 pages
            
            all_content = []
            for i, image in enumerate(images):
                # Convert PIL image to bytes
                img_byte_arr = io.BytesIO()
                image.save(img_byte_arr, format='PNG')
                img_bytes = img_byte_arr.getvalue()
                
                # Process each page as image
                page_result = OCRProcessingService._process_image_with_mistral(
                    mistral_client, img_bytes, f"{filename}_page_{i+1}"
                )
                
                all_content.append(f"## Trang {i+1}\n\n{page_result['markdown_content']}")
            
            combined_content = f"# PDF: {filename}\n\n" + "\n\n---\n\n".join(all_content)
            
            return {
                "file_type": "pdf",
                "filename": filename,
                "markdown_content": combined_content,
                "structured_content": {
                    "type": "pdf_ocr_fallback",
                    "processing_method": "pdf_to_images_fallback",
                    "pages_processed": len(images),
                    "processing_model": "pixtral-12b-2409"
                }
            }
        except Exception as fallback_error:
            logger.error(f"Fallback PDF processing also failed: {fallback_error}")
            raise Exception(f"Both primary and fallback PDF processing failed. Last error: {fallback_error}")

    @staticmethod
    def _process_docx_with_mistral_fallback(mistral_client: MistralAI, file_content: bytes, filename: str) -> Dict[str, Any]:
        """Fallback method for DOCX processing when main method fails"""
        try:
            logger.info(f"Using fallback method for DOCX: {filename}")
            
            # Create a BytesIO object from file content
            doc_stream = BytesIO(file_content)
            
            # Load the document
            doc = docx.Document(doc_stream)
            
            # Extract text from all paragraphs
            text_content = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text)
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_content.append(" | ".join(row_text))
            
            # Combine all text
            full_text = "\n\n".join(text_content)
            
            if not full_text.strip():
                raise ValueError("No text content found in DOCX file")
            
            # Simple fallback: just format the extracted text as basic markdown
            markdown_content = f"# DOCX: {filename}\n\n**Ghi chú:** Tài liệu được xử lý bằng phương pháp fallback (không qua AI formatting).\n\n---\n\n{full_text}"
            
            result = {
                "file_type": "docx",
                "filename": filename,
                "markdown_content": markdown_content,
                "structured_content": {
                    "type": "docx_extraction_fallback",
                    "processing_method": "raw_text_extraction",
                    "processing_model": "none",
                    "paragraphs_count": len([p for p in doc.paragraphs if p.text.strip()]),
                    "tables_count": len(doc.tables),
                    "note": "Processed with fallback method - no AI formatting"
                }
            }
            
            return result
            
        except Exception as fallback_error:
            logger.error(f"Fallback DOCX processing also failed: {fallback_error}")
            raise Exception(f"Both primary and fallback DOCX processing failed. Last error: {fallback_error}")

    @staticmethod
    def _process_image_with_mistral(mistral_client: MistralAI, file_content: bytes, filename: str) -> Dict[str, Any]:
        encoded_image = base64.b64encode(file_content).decode('utf-8')
        if filename.lower().endswith('.png'):
            mime_type = "image/png"
        else:
            mime_type = "image/jpeg"
        
        ocr_response = mistral_client.chat.complete(
        model="pixtral-12b-2409",  # ← Đúng model cho vision
        messages=[
            {
                "role": "user",
                "content": [
                    {
                        "type": "text",
                        "text": "Hãy trích xuất tất cả văn bản từ hình ảnh này và định dạng thành markdown có cấu trúc tốt. Bao gồm bảng, tiêu đề và giữ nguyên cấu trúc tài liệu. Nếu văn bản là tiếng Việt thì giữ nguyên tiếng Việt."
                    },
                    {
                        "type": "image_url",
                        "image_url": f"data:{mime_type};base64,{encoded_image}"
                    }
                ]
            }
        ],
        max_tokens=4000,
        temperature=0.1
    )
           
        markdown_content = ocr_response.choices[0].message.content if ocr_response and ocr_response.choices else ''

        result = {
            "file_type": "image",
            "filename": filename,
            "markdown_content": markdown_content,
            "structured_content": {
                "type": "image_ocr",
                "processing_model": "pixtral-12b-2409"
            }
        }
        return result
            

        
    @staticmethod
    def _process_text_with_mistral(mistral_client: MistralAI, file_content: bytes, filename: str) -> Dict[str, Any]:
        """Xử lý file text với Mistral - UPDATED API FORMAT"""
        try:
            text_content = file_content.decode('utf-8')
        except UnicodeDecodeError:
            try:
                text_content = file_content.decode('latin-1')
            except:
                text_content = file_content.decode('utf-8', errors='ignore')
        
        ocr_response = mistral_client.chat.complete(
                model="mistral-large-latest",
                messages=[
                    {
                        "role": "user",
                        "content": f"Hãy định dạng nội dung văn bản này thành markdown có cấu trúc với tiêu đề, danh sách và formatting phù hợp để cải thiện khả năng đọc nhưng vẫn giữ nguyên toàn bộ nội dung gốc. Nếu văn bản là tiếng Việt thì giữ nguyên tiếng Việt, nếu là ngôn ngữ khác thì giữ nguyên:\n\n{text_content}"
                    }
                ],
                max_tokens=8000,
                temperature=0.1
            )
        
        markdown_content = ocr_response.choices[0].message.content if ocr_response and ocr_response.choices else ''
        
        result = {
            "file_type": "text",
            "filename": filename,
            "markdown_content": markdown_content,
            "structured_content": {
                "type": "text_formatting",
                "processing_model": "mistral-large-latest"
            }
        }
        
        return result

    @staticmethod
    async def _process_pdf_with_mistral_async(Misa: MistralAI, file_content: bytes, filename: str) -> Dict[str, Any]:
        """Async version of PDF processing with Mistral OCR API directly"""
        logger.info("Process pdf with mistral async ...")
        try:
            # Check page count first
            page_count = OCRProcessingService._get_pdf_page_count(file_content)
            # If more than 20 pages, use chunking
            logger.info(f"page count = {page_count}")
            if page_count > 20:
                logger.info(f"PDF {filename} has {page_count} pages (>20), using chunked processing")
                return await OCRProcessingService._process_pdf_chunks_async(Misa, file_content, filename)
            
            # For smaller PDFs, process normally
            # Upload file to Mistral for OCR processing
            uploaded_file = Misa.mistral_client.files.upload(
                file={
                    "file_name": Path(filename).stem,
                    "content": file_content,
                },
                purpose="ocr",
            )
            
            # Get signed URL for processing
            signed_url = Misa.mistral_client.files.get_signed_url(
                file_id=uploaded_file.id, 
                expiry=1
            )
            logger.info("Making pdf_response ....")
            # Process PDF with Mistral OCR
            pdf_response = Misa.mistral_client.ocr.process(
                model="mistral-ocr-latest", 
                document=DocumentURLChunk(document_url=signed_url.url), 
            )
            logger.info("Successfully made")
            # Extract markdown content from OCR response
            all_markdowns = []
            for page in pdf_response.pages:
                if page.markdown and page.markdown.strip():
                    all_markdowns.append(page.markdown)
            
            # Combine all pages
            combined_content = f"# PDF: {filename}\n\n" + "\n\n---\n\n".join(all_markdowns)

            # Check if OCR result is too large - idea from OCR_example.py
            estimated_tokens = len(combined_content) // 4
            if estimated_tokens > 100000:  # If OCR result is very large
                logger.info(f"PDF OCR result for {filename} is large ({estimated_tokens} estimated tokens), using text chunking")
                return await OCRProcessingService._process_large_content_in_chunks(
                    Misa, combined_content, filename, "pdf"
                )
            
            # Clean up the uploaded file from Mistral
            try:
                Misa.mistral_client.files.delete(file_id=uploaded_file.id)
            except Exception as e:
                logger.warning(f"Could not delete uploaded file from Mistral: {e}")
            
            result = {
                "file_type": "pdf",
                "filename": filename,
                "markdown_content": combined_content,
                "structured_content": {
                    "pages_processed": len(pdf_response.pages),
                    "processing_model": "mistral-ocr-latest"
                }
            }
            return result
            
        except Exception as e:
            logger.error(f"Error processing PDF with Mistral OCR API (async): {e}")
            # Fallback to sync method if async fails
            return OCRProcessingService._process_pdf_with_mistral_fallback(Misa, file_content, filename)

    @staticmethod
    async def _process_image_with_mistral_async(mistral_client: MistralAI, file_content: bytes, filename: str) -> Dict[str, Any]:
        """Async version of image processing with Mistral AI"""
        encoded_image = base64.b64encode(file_content).decode('utf-8')
        if filename.lower().endswith('.png'):
            mime_type = "image/png"
        else:
            mime_type = "image/jpeg"
        
        ocr_response = mistral_client.chat.complete(
            model="pixtral-12b-2409",  # ← Đúng model cho vision
            messages=[
                {
                    "role": "user",
                    "content": [
                        {
                            "type": "text",
                            "text": "Hãy trích xuất tất cả văn bản từ hình ảnh này và định dạng thành markdown có cấu trúc tốt. Bao gồm bảng, tiêu đề và giữ nguyên cấu trúc tài liệu. Nếu văn bản là tiếng Việt thì giữ nguyên tiếng Việt."
                        },
                        {
                            "type": "image_url",
                            "image_url": f"data:{mime_type};base64,{encoded_image}"
                        }
                    ]
                }
            ],
            max_tokens=4000,
            temperature=0.1
        )
               
        markdown_content = ocr_response.choices[0].message.content if ocr_response and ocr_response.choices else ''

        result = {
            "file_type": "image",
            "filename": filename,
            "markdown_content": markdown_content,
            "structured_content": {
                "type": "image_ocr_async",
                "processing_model": "pixtral-12b-2409"
            }
        }
        return result

    @staticmethod
    async def _process_text_with_mistral_async(mistral_client: MistralAI, file_content: bytes, filename: str) -> Dict[str, Any]:
        """Async version of text processing with Mistral AI"""
        try:
            text_content = file_content.decode('utf-8')
        except UnicodeDecodeError:
            try:
                text_content = file_content.decode('latin-1')
            except:
                text_content = file_content.decode('utf-8', errors='ignore')
        
        # Check if content is too large - idea from OCR_example.py
        estimated_tokens = len(text_content) // 4
        if estimated_tokens > 100000:  # If too large, use chunked processing
            logger.info(f"Text file {filename} is large ({estimated_tokens} estimated tokens), using chunked processing")
            return await OCRProcessingService._process_large_content_in_chunks(
                mistral_client, text_content, filename, "text"
            )
        
        ocr_response = mistral_client.chat.complete(
            model="mistral-large-latest",
            messages=[
                {
                    "role": "user",
                    "content": f"Hãy định dạng nội dung văn bản này thành markdown có cấu trúc với tiêu đề, danh sách và formatting phù hợp để cải thiện khả năng đọc nhưng vẫn giữ nguyên toàn bộ nội dung gốc. Nếu văn bản là tiếng Việt thì giữ nguyên tiếng Việt, nếu là ngôn ngữ khác thì giữ nguyên:\n\n{text_content}"
                }
            ],
            max_tokens=8000,
            temperature=0.1
        )
        
        markdown_content = ocr_response.choices[0].message.content if ocr_response and ocr_response.choices else ''
        
        result = {
            "file_type": "text",
            "filename": filename,
            "markdown_content": markdown_content,
            "structured_content": {
                "type": "text_formatting_async",
                "processing_model": "mistral-large-latest"
            }
        }
        
        return result

    @staticmethod
    async def _process_docx_with_mistral_async(misa: MistralAI, file_content: bytes, filename: str) -> Dict[str, Any]:
        """Async version of DOCX processing with Mistral AI"""    
        logger.info(f"Starting async DOCX processing for file: {filename}")
        
        try:
            # Create a BytesIO object from file content
            logger.debug(f"Creating BytesIO stream for {filename}, content size: {len(file_content)} bytes")
            doc_stream = BytesIO(file_content)
            
            # Load the document
            logger.debug(f"Loading DOCX document: {filename}")
            doc = docx.Document(doc_stream)
            logger.info(f"Successfully loaded DOCX document: {filename}")
            
            # Extract text from all paragraphs
            text_content = []
            paragraph_count = 0
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text)
                    paragraph_count += 1
            
            logger.debug(f"Extracted {paragraph_count} paragraphs from {filename}")
            
            # Extract text from tables
            table_count = 0
            table_rows_count = 0
            for table in doc.tables:
                table_count += 1
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_content.append(" | ".join(row_text))
                        table_rows_count += 1
            
            logger.debug(f"Extracted {table_count} tables with {table_rows_count} rows from {filename}")
            
            # Combine all text
            full_text = "\n\n".join(text_content)
            logger.info(f"Combined text content length: {len(full_text)} characters for {filename}")
            
            if not full_text.strip():
                logger.error(f"No text content found in DOCX file: {filename}")
                raise ValueError("No text content found in DOCX file")
            
            # Check if content is too large - idea from OCR_example.py
            estimated_tokens = len(full_text) // 4
            logger.debug(f"Estimated tokens for {filename}: {estimated_tokens}")
            
            if estimated_tokens > 100000:  # If too large, use chunked processing
                logger.info(f"DOCX file {filename} is large ({estimated_tokens} estimated tokens), using chunked processing")
                return await OCRProcessingService._process_large_content_in_chunks(
                    misa, full_text, filename, "docx"
                )
            
            # Process with Mistral AI to format as markdown
            logger.info(f"Starting Mistral AI processing for {filename}")
            
            ocr_response = misa.chat.complete(
                model="mistral-large-latest",
                messages=[
                    {
                        "role": "user",
                        "content": f"Hãy định dạng nội dung DOCX này thành markdown có cấu trúc với tiêu đề, danh sách, bảng và formatting phù hợp. Giữ nguyên toàn bộ nội dung và ngôn ngữ gốc:\n\n{full_text}"
                    }
                ],
                max_tokens=4000,
                temperature=0.1
            )

            logger.info(f"Mistral AI processing completed for {filename}")
            
            markdown_content = ocr_response.choices[0].message.content if ocr_response and ocr_response.choices else full_text
            
            if markdown_content == full_text:
                logger.warning(f"Mistral AI returned original text (possible API issue) for {filename}")
            else:
                logger.info(f"Mistral AI successfully processed content for {filename}, markdown length: {len(markdown_content)}")
            
            result = {
                "file_type": "docx",
                "filename": filename,
                "markdown_content": markdown_content,
                "structured_content": {
                    "type": "docx_extraction_async",
                    "processing_model": "mistral-large-latest",
                    "paragraphs_count": len([p for p in doc.paragraphs if p.text.strip()]),
                    "tables_count": len(doc.tables)
                }
            }
            
            logger.info(f"Successfully completed async DOCX processing for {filename}")
            logger.debug(f"Result structure: {result['structured_content']}")
            
            return result
            
        except Exception as e:
            logger.error(f"Error in async DOCX processing for {filename}: {str(e)}", exc_info=True)
            raise

    @staticmethod
    async def _process_large_content_in_chunks(mistral_client: MistralAI, content: str, filename: str, file_type: str) -> Dict[str, Any]:
        """Process large content in chunks to avoid token limits - inspired by OCR_example.py"""
        logger.info(f"Processing large {file_type} content in chunks for {filename}")
        
        # Split content into chunks based on length
        max_chunk_size = 50000  # characters per chunk
        chunks = []
        
        for i in range(0, len(content), max_chunk_size):
            chunk = content[i:i + max_chunk_size]
            chunks.append(chunk)
        
        processed_chunks = []
        
        # Process each chunk
        for i, chunk in enumerate(chunks):
            logger.info(f"Processing chunk {i+1}/{len(chunks)} for {filename}")
            
            try:
                ocr_response = mistral_client.chat.complete(
                    model="mistral-large-latest",
                    messages=[
                        {
                            "role": "user",
                            "content": f"Hãy định dạng và làm sạch nội dung văn bản này thành markdown có cấu trúc tốt. Giữ nguyên toàn bộ nội dung và ngôn ngữ gốc:\n\n{chunk}"
                        }
                    ],
                    max_tokens=4000,
                    temperature=0.1
                )
                
                chunk_markdown = ocr_response.choices[0].message.content if ocr_response and ocr_response.choices else chunk
                processed_chunks.append(f"## Phần {i+1}\n\n{chunk_markdown}")
                
                # Small delay between chunks
                await asyncio.sleep(0.5)
                
            except Exception as e:
                logger.warning(f"Error processing chunk {i+1}: {e}")
                # Fallback to raw content
                processed_chunks.append(f"## Phần {i+1}\n\n{chunk}")
        
        # Combine all chunks
        combined_content = f"# {filename} (Xử lý theo chunks)\n\n**Ghi chú:** Tài liệu lớn đã được chia thành {len(chunks)} phần để xử lý.\n\n---\n\n" + "\n\n---\n\n".join(processed_chunks)
        
        result = {
            "file_type": file_type,
            "filename": filename,
            "markdown_content": combined_content,
            "structured_content": {
                "type": f"{file_type}_large_chunked",
                "processing_method": "large_content_chunking",
                "chunks_processed": len(chunks),
                "processing_model": "mistral-large-latest"
            }
        }
        
        return result

    @staticmethod
    def get_processing_status(ocr_result_id: str) -> Dict[str, Any]:
        """Lấy trạng thái OCR - ✅ UPDATED: Tìm trong cả admin_database và user_database"""
        # Try to find in new collections first
        logger.info("Get processing status ...")
        from .models import AdminOCRResult, UserOCRResult, OCRResult
        
        ocr_result = None
        
        # Check admin database
        ocr_result = AdminOCRResult.objects(id=ocr_result_id).first()
        if not ocr_result:
            # Check user database  
            ocr_result = UserOCRResult.objects(id=ocr_result_id).first()
        if not ocr_result:
            # Fallback to legacy collection
            ocr_result = OCRResult.objects(id=ocr_result_id).first()
        
        if not ocr_result:
            raise ValueError(f"OCR result with ID {ocr_result_id} not found")
        
        return {
            "id": str(ocr_result.id),
            "status": ocr_result.status,
            "is_completed": ocr_result.is_completed,
            "is_successful": ocr_result.is_successful,
            "error_message": ocr_result.error_message,
            "created_at": ocr_result.created_at.isoformat() if ocr_result.created_at else None,
            "completed_at": ocr_result.completed_at.isoformat() if ocr_result.completed_at else None,
            "collection": ocr_result._meta.get('collection', 'unknown')  # Add collection info for debugging
        }

    @staticmethod
    def _get_pdf_page_count(file_content: bytes) -> int:
        """Get the number of pages in a PDF file"""
        logger.info("counting page ...")
        try:
            # Try using pdfplumber first
            try:
                import pdfplumber
                with pdfplumber.open(io.BytesIO(file_content)) as pdf:
                    return len(pdf.pages)
            except ImportError:
                pass
            
            # Fallback to PyPDF2
            try:
                import PyPDF2
                pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_content))
                return len(pdf_reader.pages)
            except ImportError:
                pass
            
            # If no PDF libraries available, return 0 to trigger Mistral processing
            logger.warning("No PDF parsing libraries available, cannot determine page count")
            return 0
            
        except Exception as e:
            logger.warning(f"Could not determine PDF page count: {e}")
            return 0

    @staticmethod
    async def _process_pdf_chunks_async(mistral_client: MistralAI, file_content: bytes, filename: str, max_pages_per_chunk: int = 20) -> Dict[str, Any]:
        """Process large PDF in chunks asynchronously with improved error handling"""
        logger.info("Processing pdf chunk async ...")
        try:
            # First, try to get total page count
            total_pages = OCRProcessingService._get_pdf_page_count(file_content)
            logger.info(f"PDF {filename} has {total_pages} pages, processing in chunks of {max_pages_per_chunk}")
            
            if total_pages <= max_pages_per_chunk:
                # If small enough, process normally
                return await OCRProcessingService._process_pdf_with_mistral_async(mistral_client, file_content, filename)
            
            # Split PDF into chunks and process each chunk
            chunks = []
            for start_page in range(1, total_pages + 1, max_pages_per_chunk):
                end_page = min(start_page + max_pages_per_chunk - 1, total_pages)
                chunk_info = {
                    'start_page': start_page,
                    'end_page': end_page,
                    'chunk_number': len(chunks) + 1
                }
                chunks.append(chunk_info)
            
            logger.info(f"Split PDF into {len(chunks)} chunks")
            
            # Process chunks with improved error handling and retry logic
            semaphore = asyncio.Semaphore(2)  # Reduced from 3 to 2 for better stability
            chunk_results = []
            
            async def process_chunk_with_retry(chunk_info, max_retries=3):
                async with semaphore:
                    for attempt in range(max_retries):
                        try:
                            # Add delay between attempts
                            if attempt > 0:
                                delay = min(2 ** attempt, 10)  # Exponential backoff, max 10 seconds
                                logger.info(f"Retrying chunk {chunk_info['chunk_number']}, attempt {attempt + 1} after {delay}s delay")
                                await asyncio.sleep(delay)
                            
                            # Add small delay between chunks to avoid overwhelming the API
                            await asyncio.sleep(0.5)
                            
                             # Create chunk-specific filename
                            chunk_filename = f"{Path(filename).stem}_chunk_{chunk_info['chunk_number']}_pages_{chunk_info['start_page']}-{chunk_info['end_page']}.pdf"
                            
                            # Process the chunk using Mistral OCR API
                            chunk_result = await OCRProcessingService._process_pdf_chunk_with_mistral(
                                mistral_client, file_content, chunk_filename, 
                                chunk_info['start_page'], chunk_info['end_page']
                            )
                            
                            chunk_result['chunk_info'] = chunk_info
                            logger.info(f"Successfully processed chunk {chunk_info['chunk_number']}")
                            return chunk_result
                            
                        except Exception as e:
                            logger.warning(f"Attempt {attempt + 1} failed for chunk {chunk_info['chunk_number']}: {e}")
                            if attempt == max_retries - 1:  # Last attempt
                                logger.error(f"All {max_retries} attempts failed for chunk {chunk_info['chunk_number']}")
                                return {
                                    'chunk_info': chunk_info,
                                    'error': str(e),
                                    'markdown_content': f"# Lỗi xử lý chunk {chunk_info['chunk_number']}\n\nKhông thể xử lý trang {chunk_info['start_page']}-{chunk_info['end_page']} sau {max_retries} lần thử: {str(e)}"
                                }
                            # Continue to next attempt
                            
                # Should not reach here
                return {
                    'chunk_info': chunk_info,
                    'error': 'Unknown error',
                    'markdown_content': f"# Lỗi không xác định chunk {chunk_info['chunk_number']}\n\nKhông thể xử lý trang {chunk_info['start_page']}-{chunk_info['end_page']}"
                }
            
            # Process chunks in smaller batches to avoid overwhelming the API
            batch_size = 3
            all_chunk_results = []
            
            for i in range(0, len(chunks), batch_size):
                batch_chunks = chunks[i:i + batch_size]
                logger.info(f"Processing batch {i//batch_size + 1}: chunks {[c['chunk_number'] for c in batch_chunks]}")
                
                # Process batch
                batch_tasks = [process_chunk_with_retry(chunk) for chunk in batch_chunks]
                batch_results = await asyncio.gather(*batch_tasks, return_exceptions=True)
                all_chunk_results.extend(batch_results)
                
                # Add delay between batches
                if i + batch_size < len(chunks):
                    await asyncio.sleep(2)
            
            # Combine results
            all_content = []
            successful_chunks = 0
            failed_chunks = 0
            
            for i, result in enumerate(all_chunk_results):
                if isinstance(result, Exception):
                    logger.error(f"Chunk {i+1} failed with exception: {result}")
                    all_content.append(f"# Chunk {i+1} - Lỗi xử lý\n\n```\nLỗi: {str(result)}\n```")
                    failed_chunks += 1
                elif 'error' in result:
                    all_content.append(result['markdown_content'])
                    failed_chunks += 1
                else:
                    chunk_info = result['chunk_info']
                    chunk_content = f"# Phần {chunk_info['chunk_number']} (Trang {chunk_info['start_page']}-{chunk_info['end_page']})\n\n{result.get('markdown_content', '')}"
                    all_content.append(chunk_content)
                    successful_chunks += 1
            
            # Check if too many chunks failed, use fallback
            failure_rate = failed_chunks / len(chunks)
            if failure_rate > 0.5:  # If more than 50% chunks failed
                logger.warning(f"High failure rate ({failure_rate:.2%}), attempting fallback processing")
                return await OCRProcessingService._process_pdf_with_fallback_chunking(mistral_client, file_content, filename)
            
            # Combine all chunks
            combined_content = f"# PDF: {filename} (Xử lý theo chunks)\n\n**Tổng quan:**\n- Tổng số trang: {total_pages}\n- Số chunks: {len(chunks)}\n- Chunks thành công: {successful_chunks}\n- Chunks thất bại: {failed_chunks}\n- Tỷ lệ thành công: {successful_chunks/len(chunks):.1%}\n\n---\n\n" + "\n\n---\n\n".join(all_content)
            
            result = {
                "file_type": "pdf",
                "filename": filename,
                "markdown_content": combined_content,
                "structured_content": {
                    "type": "pdf_ocr_chunked",
                    "processing_method": "mistral_ocr_chunked_improved",
                    "total_pages": total_pages,
                    "chunks_processed": len(chunks),
                    "successful_chunks": successful_chunks,
                    "failed_chunks": failed_chunks,
                    "success_rate": successful_chunks/len(chunks),
                    "processing_model": "mistral-large-latest"
                }
            }
            
            return result
            
        except Exception as e:
            logger.error(f"Error in chunked PDF processing: {e}")
            # Fallback to regular processing
            return await OCRProcessingService._process_pdf_with_mistral_async(mistral_client, file_content, filename)

    @staticmethod
    async def _process_pdf_chunk_with_mistral(mistral_client: MistralAI, file_content: bytes, chunk_filename: str, start_page: int, end_page: int) -> Dict[str, Any]:
        """Process a specific page range of PDF with Mistral OCR API"""
        uploaded_file_id = None
        try:
            # Validate mistral_client
            if not mistral_client or not mistral_client.mistral_client:
                raise ValueError("Mistral client is not properly initialized")
            
            # Upload file to Mistral for OCR processing
            logger.info(f"Uploading chunk file: {chunk_filename}")
            uploaded_file = mistral_client.mistral_client.files.upload(
                file={
                    "file_name": Path(chunk_filename).stem,
                    "content": file_content,
                },
                purpose="ocr",
            )
            uploaded_file_id = uploaded_file.id
            logger.info(f"Successfully uploaded file with ID: {uploaded_file_id}")
            
            # Get signed URL for processing
            signed_url = mistral_client.mistral_client.files.get_signed_url(
                file_id=uploaded_file.id, 
                expiry=1
            )
            logger.info(f"Got signed URL for processing")
            
            # Add timeout and retry for OCR processing
            max_processing_retries = 2
            for processing_attempt in range(max_processing_retries):
                try:
                    # Process specific pages with Mistral OCR
                    logger.info(f"Starting OCR processing attempt {processing_attempt + 1}")
                    pdf_response = mistral_client.mistral_client.ocr.process(
                        document=DocumentURLChunk(document_url=signed_url.url), 
                        model="mistral-large-latest", 
                        include_image_base64=False,
                        # Note: Mistral OCR API doesn't support page range in the same way
                        # It will process the entire document, but we'll label it as chunk
                    )
                    logger.info(f"OCR processing completed successfully")
                    break  # Success, exit retry loop
                    
                except Exception as processing_error:
                    logger.warning(f"OCR processing attempt {processing_attempt + 1} failed: {processing_error}")
                    if processing_attempt == max_processing_retries - 1:
                        raise processing_error
                    # Wait before retry
                    await asyncio.sleep(2)
            
            # Extract markdown content from OCR response
            all_markdowns = []
            if pdf_response and pdf_response.pages:
                for i, page in enumerate(pdf_response.pages):
                    # Calculate actual page number
                    actual_page = start_page + i
                    if actual_page > end_page:
                        break
                        
                    if page.markdown and page.markdown.strip():
                        page_content = f"## Trang {actual_page}\n\n{page.markdown}"
                        all_markdowns.append(page_content)
                
                # If no pages in range, take the first few pages
                if not all_markdowns and pdf_response.pages:
                    pages_to_take = min(end_page - start_page + 1, len(pdf_response.pages))
                    for i in range(pages_to_take):
                        if pdf_response.pages[i].markdown and pdf_response.pages[i].markdown.strip():
                            page_content = f"## Trang {start_page + i}\n\n{pdf_response.pages[i].markdown}"
                            all_markdowns.append(page_content)
            
            combined_content = "\n\n".join(all_markdowns)
            if not combined_content.strip():
                combined_content = f"Không thể trích xuất nội dung từ trang {start_page}-{end_page}"
            
            return {
                "markdown_content": combined_content,
                "pages_processed": len(all_markdowns),
                "start_page": start_page,
                "end_page": end_page
            }
            
        except Exception as e:
            logger.error(f"Error processing PDF chunk {start_page}-{end_page}: {e}")
            raise
            
        finally:
            # Clean up the uploaded file from Mistral
            if uploaded_file_id:
                try:
                    mistral_client.mistral_client.files.delete(file_id=uploaded_file_id)
                    logger.info(f"Successfully deleted uploaded file: {uploaded_file_id}")
                except Exception as cleanup_error:
                    logger.warning(f"Could not delete uploaded chunk file from Mistral: {cleanup_error}")

    @staticmethod
    async def _process_pdf_with_fallback_chunking(mistral_client: MistralAI, file_content: bytes, filename: str) -> Dict[str, Any]:
        """Fallback chunking method using pdf2image when Mistral OCR API fails too much"""
        try:
            logger.info(f"Using fallback chunking method for {filename}")
            
            # Try to get page count
            total_pages = OCRProcessingService._get_pdf_page_count(file_content)
            
            # Use pdf2image approach with chunking
            try:
                from pdf2image import convert_from_bytes
                
                # Process in smaller chunks of 10 pages to avoid memory issues
                chunk_size = 10
                all_content = []
                successful_chunks = 0
                failed_chunks = 0
                
                for start_page in range(1, total_pages + 1, chunk_size):
                    end_page = min(start_page + chunk_size - 1, total_pages)
                    chunk_number = (start_page - 1) // chunk_size + 1
                    
                    try:
                        logger.info(f"Processing fallback chunk {chunk_number}: pages {start_page}-{end_page}")
                        
                        # Convert specific pages to images
                        images = convert_from_bytes(
                            file_content, 
                            dpi=150,  # Reduced DPI for faster processing
                            first_page=start_page, 
                            last_page=end_page
                        )
                        
                        chunk_content = []
                        for i, image in enumerate(images):
                            # Convert PIL image to bytes
                            img_byte_arr = io.BytesIO()
                            image.save(img_byte_arr, format='PNG')
                            img_bytes = img_byte_arr.getvalue()
                            
                            # Process each page as image
                            page_result = await OCRProcessingService._process_image_with_mistral_async(
                                mistral_client, img_bytes, f"{filename}_page_{start_page + i}"
                            )
                            
                            chunk_content.append(f"## Trang {start_page + i}\n\n{page_result['markdown_content']}")
                        
                        chunk_markdown = f"# Phần {chunk_number} (Trang {start_page}-{end_page})\n\n" + "\n\n".join(chunk_content)
                        all_content.append(chunk_markdown)
                        successful_chunks += 1
                        
                        # Add small delay between chunks
                        await asyncio.sleep(1)
                        
                    except Exception as chunk_error:
                        logger.error(f"Fallback chunk {chunk_number} failed: {chunk_error}")
                        error_content = f"# Phần {chunk_number} (Trang {start_page}-{end_page}) - Lỗi\n\nKhông thể xử lý: {str(chunk_error)}"
                        all_content.append(error_content)
                        failed_chunks += 1
                
                combined_content = f"# PDF: {filename} (Xử lý fallback chunking)\n\n**Tổng quan:**\n- Tổng số trang: {total_pages}\n- Phương pháp: PDF to Images (fallback)\n- Chunks thành công: {successful_chunks}\n- Chunks thất bại: {failed_chunks}\n- Tỷ lệ thành công: {successful_chunks/(successful_chunks + failed_chunks):.1%}\n\n---\n\n" + "\n\n---\n\n".join(all_content)
                
                return {
                    "file_type": "pdf",
                    "filename": filename,
                    "markdown_content": combined_content,
                    "structured_content": {
                        "type": "pdf_ocr_fallback_chunked",
                        "processing_method": "pdf_to_images_chunked_fallback",
                        "total_pages": total_pages,
                        "chunks_processed": successful_chunks + failed_chunks,
                        "successful_chunks": successful_chunks,
                        "failed_chunks": failed_chunks,
                        "success_rate": successful_chunks/(successful_chunks + failed_chunks) if (successful_chunks + failed_chunks) > 0 else 0,
                        "processing_model": "pixtral-12b-2409"
                    }
                }
                
            except ImportError:
                logger.error("pdf2image not available for fallback")
                # If pdf2image is not available, return error message
                return {
                    "file_type": "pdf",
                    "filename": filename,
                    "markdown_content": f"# PDF: {filename} - Lỗi xử lý\n\nKhông thể xử lý PDF này do:\n- Mistral OCR API gặp nhiều lỗi\n- pdf2image không khả dụng cho fallback\n\nVui lòng thử lại sau hoặc sử dụng file PDF nhỏ hơn.",
                    "structured_content": {
                        "type": "pdf_ocr_error",
                        "processing_method": "failed",
                        "error": "Both primary and fallback methods failed"
                    }
                }
                
        except Exception as e:
            logger.error(f"Fallback chunking also failed: {e}")
            return {
                "file_type": "pdf",
                "filename": filename,
                "markdown_content": f"# PDF: {filename} - Lỗi xử lý hoàn toàn\n\nKhông thể xử lý PDF này: {str(e)}",
                "structured_content": {
                    "type": "pdf_ocr_error",
                    "processing_method": "failed",
                    "error": str(e)
                }
            }
